#!/usr/bin/python
#encoding=utf-8
from googletrans import Translator
import bibtexparser
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import RGBColor
import os
import tkinter.filedialog
import tkinter.messagebox

gtrans = Translator(service_urls=[
    'translate.google.cn'
])

doc = Document()
styles = doc.styles
style_en_title = styles.add_style('EnTitle', WD_STYLE_TYPE.PARAGRAPH)
style_en_title.base_style = styles['Heading 1']
style_en_title.font.name = 'Arial'
style_en_title.font.size = Pt(12)
style_en_title.font.color.rgb = RGBColor(0, 0, 0)
style_en_title.paragraph_format.space_before = Pt(20)
style_en_title.paragraph_format.space_after = Pt(5)
style_en_title.paragraph_format.line_spacing = 1.5

style_en_author = styles.add_style('EnAuthor', WD_STYLE_TYPE.PARAGRAPH)
style_en_author.base_style = styles['Normal']
style_en_author.font.name = 'Calibri'
style_en_author.font.size = Pt(10)
style_en_author.font.color.rgb = RGBColor(0, 0, 0)
style_en_author.font.underline = True
style_en_author.paragraph_format.space_before = 0
style_en_author.paragraph_format.space_after = 0
style_en_author.paragraph_format.line_spacing = 1.5

style_en_note = styles.add_style('EnNote', WD_STYLE_TYPE.PARAGRAPH)
style_en_note.base_style = styles['Normal']
style_en_note.font.name = 'Calibri'
style_en_note.font.size = Pt(10)
style_en_note.font.color.rgb = RGBColor(0, 0, 0)
style_en_note.paragraph_format.space_before = 0
style_en_note.paragraph_format.space_after = 0
style_en_note.paragraph_format.line_spacing = 1.5

style_en_pmid = styles.add_style('EnPMID', WD_STYLE_TYPE.PARAGRAPH)
style_en_pmid.base_style = styles['Normal']
style_en_pmid.font.name = 'Calibri'
style_en_pmid.font.size = Pt(10)
style_en_pmid.font.color.rgb = RGBColor(87, 87, 87)
style_en_pmid.paragraph_format.space_before = 0
style_en_pmid.paragraph_format.space_after = 0
style_en_pmid.paragraph_format.line_spacing = 1.5

style_en_abshead = styles.add_style('EnAbshead', WD_STYLE_TYPE.PARAGRAPH)
style_en_abshead.base_style = styles['Normal']
style_en_abshead.font.name = 'Times New Roman'
style_en_abshead.font.size = Pt(11)
style_en_abshead.font.bold = True
style_en_abshead.font.color.rgb = RGBColor(0, 0, 0)
style_en_abshead.paragraph_format.space_before = Pt(10)
style_en_abshead.paragraph_format.space_after = 0
style_en_abshead.paragraph_format.line_spacing = 1.5

style_en_abstract = styles.add_style('EnAbstract', WD_STYLE_TYPE.PARAGRAPH)
style_en_abstract.base_style = styles['Normal']
style_en_abstract.font.name = 'Times New Roman'
style_en_abstract.font.size = Pt(10)
style_en_abstract.font.color.rgb = RGBColor(0, 0, 0)
style_en_abstract.paragraph_format.space_before = 0
style_en_abstract.paragraph_format.space_after = 0
style_en_abstract.paragraph_format.line_spacing = 1.5

style_en_authoraddr = styles.add_style(
    'EnAuthorAddr', WD_STYLE_TYPE.PARAGRAPH)
style_en_authoraddr.base_style = styles['EnAuthorAddr']
style_en_authoraddr.font.size = Pt(8)

style_zh_title = styles.add_style('ZhTitle', WD_STYLE_TYPE.PARAGRAPH)
style_zh_title.base_style = styles['EnTitle']
style_zh_title.font.name = u'宋体'
style_zh_title._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

style_zh_abshead = styles.add_style('ZhAbshead', WD_STYLE_TYPE.PARAGRAPH)
style_zh_abshead.base_style = styles['EnAbshead']
style_zh_abshead.font.name = u'宋体'
style_zh_abshead._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

style_zh_abstract = styles.add_style('ZhAbstract', WD_STYLE_TYPE.PARAGRAPH)
style_zh_abstract.base_style = styles['EnAbstract']
style_zh_abstract.font.name = u'宋体'
style_zh_abstract._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

style_zh_authoraddr = styles.add_style(
    'ZhAuthorAddr', WD_STYLE_TYPE.PARAGRAPH)
style_zh_authoraddr.base_style = styles['EnAuthorAddr']
style_zh_authoraddr.font.name = u'宋体'
style_zh_authoraddr._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

def do_translate(instr):
    result = gtrans.translate(instr, src='en', dest='zh-cn')
    return result.text

def do_parse_and_translate(filename):
    try:
        bibtex_file = open(filename)
    except IOError:
        tkinter.messagebox.showerror('错误', "无法读取文件%s"%filename)
        exit(1)
    bibtex_string = bibtex_file.read().encode('gbk').decode('utf-8')
    bib_db = bibtexparser.loads(bibtex_string)
    bib_entries = bib_db.entries
    print("获取到%d条记录"%len(bib_entries))
    ind = 1
    for entry in bib_entries:

        author   = entry.get('author','')
        author   = author.replace(',', '')
        author   = author.replace('. and ', ', ')
        author   = author.replace('. ', '')
        title    = entry.get('title', '')
        note     = entry.get('note',  '').split("\n").pop()
        abstract = entry.get('abstract','')
        pmid     = entry.get('accnum',  '')
        authoraddr = entry.get('authoraddr', '')
        doc.add_paragraph(style=style_en_title).add_run(title)
        doc.add_paragraph(style=style_en_author).add_run(author)
        doc.add_paragraph(style=style_en_note).add_run(note)
        if(pmid != ''):
            doc.add_paragraph(style=style_en_pmid).add_run("PMID:"+pmid)
        doc.add_paragraph(style=style_en_abshead).add_run("Abstract")
        doc.add_paragraph(style=style_en_abstract).add_run(abstract)
        if(authoraddr != ''):
            doc.add_paragraph(style=style_en_abstract).add_run("")
            doc.add_paragraph(style=style_en_authoraddr).add_run(authoraddr)

        print("正在翻译第%d/%d条记录:%s"%(ind,len(bib_entries),title))
        ind = ind + 1
        title      = do_translate(title)
        abstract   = do_translate(abstract)
        authoraddr = do_translate(authoraddr)
        doc.add_paragraph(style=style_zh_title).add_run(title)
        doc.add_paragraph(style=style_en_author).add_run(author)
        doc.add_paragraph(style=style_en_note).add_run(note)
        if(pmid != ''):
            doc.add_paragraph(style=style_en_pmid).add_run("PMID:"+pmid)
        doc.add_paragraph(style=style_zh_abshead).add_run("摘要")
        doc.add_paragraph(style=style_zh_abstract).add_run(abstract)
        if(authoraddr != ''):
            doc.add_paragraph(style=style_zh_abstract).add_run("")
            doc.add_paragraph(style=style_zh_authoraddr).add_run(authoraddr)


default_dir = r"C:\Users\Administrator\Desktop"
filenames = tkinter.filedialog.askopenfilenames(title=u"选择文件", initialdir=(os.path.expanduser(default_dir)),
filetypes=[("text file","*.txt")])
outfile = ''
if(len(filenames) == 1):
    outfile = filenames[0].replace(".txt",".docx")
    yesno = tkinter.messagebox.askyesno("问题","是否使用\'%s\'作为输出文件名"%outfile)
    if(yesno == False):
        outfile = tkinter.filedialog.asksaveasfilename(title=u"请输入文件名", initialdir=(os.path.expanduser(default_dir)),
filetypes=[("work docx file","*.docx")], defaultextension=".docx")
elif(len(filenames) > 1):
    yesno = tkinter.messagebox.askyesno("问题", "选择了多个文件，将合并为一个docx，是否继续")
    if(yesno == True):
        outfile = tkinter.filedialog.asksaveasfilename(title=u"请输入合并后的文件名称", initialdir=(os.path.expanduser(default_dir)),
filetypes=[("work docx file","*.docx")], defaultextension=".docx")
    else:
        exit()
else:
    tkinter.messagebox.showerror("错误","未选择任何文件!")
    exit()
if(outfile == ''):
    tkinter.messagebox.showerror("错误", "未指定输出文件")
    exit()

print("共选取了%d个文件"%len(filenames))
index = 1
for fname in filenames:
    print("正在解析并翻译第%d/%d个文件:%s" % (index, len(filenames), fname))
    index = index + 1
    do_parse_and_translate(fname)
while(True):
    try:
        doc.save(outfile)
        tkinter.messagebox.showinfo("提示","翻译结果已经生成在%s"%outfile)
        break
    except IOError:
        retry = tkinter.messagebox.askretrycancel("错误", "写入%s失败，请关闭文档后重试"%outfile)
        if(retry==False):
            break

